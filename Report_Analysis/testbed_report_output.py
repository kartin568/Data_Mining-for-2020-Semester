import sys
import requests
from bs4 import BeautifulSoup as bs
import re
import csv
from collections import Counter

import docx
from docx import Document
from  docx.oxml.ns import qn
from docx.shared import Pt,RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Cm, Inches

import win32com.client
'''
更改的时候注意，基本信息都要传进来
'''
#文件头
file_names = {}
tmp_names = []
tmp_length = []
tmp_comment = []

fin_comment_rate = 0
whole_length = 0

testbed_rts_filename = sys.argv[1]
rts_filename = sys.argv[2]
#基本信息
software = sys.argv[7]
version = sys.argv[6]
tester = sys.argv[5]
test_date = sys.argv[4]
test_tool = sys.argv[3]
#静态质量度量
function_count = []
function_list = []
function_flag_com = 0
function_flag_fan = 0
max_min_line = []

fan_out=[]
com = []

#代码结构分析
depth = 0
unused_p_assign_twice = False
used_before_assign = False
unused_after_assign = False
unreachable_code = False
infinite_loop = False
devide_zero = False
array_bound_exceeded = False
brackets_error = False
uncalled_function = False
formal_actual_p_same_shuxing = False
formal_actual_p_same_count = False
for_cycle = False
file = open(testbed_rts_filename, 'r', encoding='gb2312')
soup = bs(file, 'lxml')
all_table = soup.find_all("table")

for table in all_table:
    head = table.find_all("th")
    trs = table.find_all("tr")
    #找到总文件表

    if len(head) == 2 and head[0].getText().strip() == "Name" and head[1].getText().strip() == "Last Modification Date":
        td_all = table.find_all("td")
        for td in td_all:
            font = str(td.find("font")).strip("</a> </font>\nNone").split("\">")[-1]
            if font != '':
                tmp_names.append(font)


    #针对单个文件找到行数,模块数以及各模块
    if len(head)>6 and head[0].getText().strip() == "File" \
            and head[1].getText().strip() == "Total Ref." and head[2].getText().strip() == "Total" \
            and head[3].getText().strip() == "Executable" and head[4].getText().strip() == "Non-Executable" \
            and head[5].getText().strip() == "Number of" and head[6].getText().strip() == "Total"\
            and head[7].getText().strip() == "Expansion":
        #cnt = cnt+1
        tds = table.find_all("td")
        if(len(tds) == 8):
            tmp_length.append(int(tds[6].getText().strip(" (P)")))
            tmp_comment.append(int(tds[2].getText().split(" ")[1]))
            function_count.append((tds[5].getText().split(" ")[1]))


for name,leng in zip(tmp_names, tmp_length):
    file_names[name] = leng

for c,l in zip(tmp_comment,tmp_length):
    fin_comment_rate += c
    whole_length += l

fin_comment_rate = fin_comment_rate/whole_length

cnt = 0
cnt_com = 0
#找各模块的扇出和圈复杂度
for table in all_table:
    tmp_list = []
    fan_list = []
    com_list = []
    head = table.find_all("th")

    if len(head) > 4 and head[0].getText().strip() == "Procedure" \
                and head[1].getText().strip() == "Globals in" and head[2].getText().strip() == "File" \
                and head[3].getText().strip() == "Fan":
            cnt = cnt + 1
            if(cnt%2 != 0):
                tds = table.find_all("td")
                for i in range(int(function_count[function_flag_fan])):
                    func = tds[1+i*4].getText().strip()
                    fan_out_tmp = int(tds[1+i*4+3].getText().split(" ")[1])
                    tmp_list.append(func)
                    fan_list.append(fan_out_tmp)

                fan_out.append(fan_list)
                function_list.append(tmp_list)
                function_flag_fan = function_flag_fan + 1

    if len(head) > 5 and head[0].getText().strip() == "Procedure" \
        and head[2].getText().strip() == "Cyclomatic" and head[3].getText().strip() == "Essential" \
        and head[4].getText().strip() == "Ess. Cycl." and head[5].getText().strip() == "Structured":
            cnt_com += 1
            if (cnt_com % 2 != 0):

                tds = table.find_all("td")
                for i in range(int(function_count[function_flag_com])):
                    tmp_com = tds[3+6*i].getText().split(" ")[1]
                    com_list.append(tmp_com)

                com.append(com_list)
                function_flag_com = function_flag_com + 1

#print(function_list)
#print(fan_out)
#代码结构分析


file = open(rts_filename, 'r', encoding='gb2312')
soup = bs(file, 'lxml')
all_table = soup.find_all("table")

LDRA = []
for table in all_table:
    head = table.find_all("th")

    if len(head) == 4 and head[0].getText().strip() == "Number of Violations" \
        and head[1].getText().strip() == "LDRA Code" and head[2].getText().strip() == "Mandatory Standards" \
        and head[3].getText().strip() == "GJB_8114 Code":
        rows = table.find_all("tr")
        for row in rows:
            cols = row.find_all("td")
            if(len(cols)!=0):
                error_code = cols[1].getText().strip(' ')
                LDRA.append(error_code.replace(" ",""))

if '1J' in LDRA: #不可达语句
    unreachable_code = True
if '80X' in LDRA:#除零操作
    devide_zero = True
if '1D' in LDRA or '15D' in LDRA:
    unused_after_assign = True
if '45D' in LDRA or '123D' in LDRA or '128D' in LDRA or '129D' in LDRA or '135D' in LDRA or '136D' in LDRA:
    used_before_assign = True
if '5C' in LDRA or '28D' in LDRA:
    infinite_loop = True
if '47S' in LDRA or '64X' in LDRA or '69X' in LDRA:
    array_bound_exceeded = True
if '49S' in LDRA or '78S' in LDRA or '361S' in LDRA or '228S' in LDRA or '229S' in LDRA:
    brackets_error = True
if '98S' in LDRA or '458S' in LDRA:
    formal_actual_p_same_shuxing = True


for table in all_table:
    trs = table.find_all("tr")
    if len(trs)==6 and trs[0].find_all("td")[0].getText().strip() == "Number of procedures:":
        cd = trs[1].find_all("td")[1].getText().strip()
        if( int(cd) > 0 ):
            uncalled_function = True
        break

#从这里开始，是输出docx文件
doc = Document()

doc.styles['Normal'].font.name = u'宋体'
doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'),u'宋体')
doc.styles['Normal'].font.size = Pt(10.5)
doc.styles['Normal'].font.color.rgb = RGBColor(0,0,0)
#标题
Head = doc.add_heading("",level=1)# 这里不填标题内容
run  = Head.add_run("静态分析问题报告单")
run.font.name=u'Cambria'
run._element.rPr.rFonts.set(qn('w:eastAsia'), u'Cambria')
run.font.color.rgb = RGBColor(0,0,0)
#表头
lent = 0

for i in function_list:
    lent = lent + len(i)

print(tmp_names)
print(function_count)
tab = doc.add_table(rows=21+2+len(file_names)+lent+1,cols=6,style='Table Grid')
tab.alignment = WD_TABLE_ALIGNMENT.CENTER

#单元格合并
tab.cell(0,1).merge(tab.cell(0,2))
tab.cell(0,4).merge(tab.cell(0,5))
tab.cell(1,1).merge(tab.cell(1,2))
tab.cell(1,4).merge(tab.cell(1,5))
tab.cell(2,0).merge(tab.cell(2,2))
tab.cell(2,3).merge(tab.cell(2,5))
tab.cell(3,0).merge(tab.cell(3,5))


tab.cell(4,0).merge(tab.cell(4,5))
tab.cell(5,1).merge(tab.cell(5,2))
tab.cell(5,4).merge(tab.cell(5,5))
for i in range(len(file_names)):
    tab.cell(i+6, 1).merge(tab.cell(i+6, 2))
    tab.cell(i+6, 4).merge(tab.cell(i+6, 5))
line = 6+len(file_names)
tab.cell(line,0).merge(tab.cell(line,5))
tab.cell(line+1,1).merge(tab.cell(line+1,2))
tab.cell(line+1,4).merge(tab.cell(line+1,5))
for p in range(lent):
    tab.cell(line+2+p, 1).merge(tab.cell(line+2+p, 2))
    tab.cell(line + 2 + p, 4).merge(tab.cell(line + 2 + p, 5))

tab.cell(line+2+lent,0).merge(tab.cell(line+2+lent,5))
line = line+2+lent
tab.cell(line+1,0).merge(tab.cell(line+1,1))
tab.cell(line+1,2).merge(tab.cell(line+1,5))
for q in range(3):
    tab.cell(line+2+q, 1).merge(tab.cell(line+2+q, 2))
    tab.cell(line +2+ q, 3).merge(tab.cell(line +2+ q, 5))
tab.cell(line+2,0).merge(tab.cell(line+4,0))
line = line+4
for q in range(4):
    tab.cell(line+1+q, 1).merge(tab.cell(line+1+q, 2))
    tab.cell(line +1+ q, 3).merge(tab.cell(line +1+ q, 5))
tab.cell(line+1,0).merge(tab.cell(line+4,0))
line = line + 4
for q in range(2):
    tab.cell(line+1+q, 1).merge(tab.cell(line+1+q, 2))
    tab.cell(line +1+ q, 3).merge(tab.cell(line +1+ q, 5))
tab.cell(line+1,0).merge(tab.cell(line+2,0))
line = line+2
for q in range(3):
    tab.cell(line+1+q, 1).merge(tab.cell(line+1+q, 2))
    tab.cell(line +1+ q, 3).merge(tab.cell(line +1+ q, 5))
tab.cell(line+1,0).merge(tab.cell(line+3,0))
line = line+3
tab.cell(line+1,0).merge(tab.cell(line+1,5))
tab.cell(line+2,1).merge(tab.cell(line+2,5))
#填充内容

tab.cell(0,0).text = "软件名称"
tab.cell(0,1).text = software
tab.cell(0,3).text = "被测件版本标识"
tab.cell(0,4).text = version

tab.cell(1,0).text = "测试人员"
tab.cell(1,1).text = tester
tab.cell(1,3).text = "测试日期"
tab.cell(1,4).text = test_date

tab.cell(2,0).text = "软件测试工具"
tab.cell(2,3).text = test_tool

c30 = "\n  软件规模总行数："+str(whole_length)+"行, 总注释率："+ str(round(fin_comment_rate*100,2))+" %\n  各文件代码行数：\n"
for c in file_names.keys():
    c30 = c30+"  "+c+":\t"+str(file_names[c])+"\n"
tab.cell(3,0).text = c30
tab.cell(4,0).text = "静态质量度量"
tab.cell(5,0).text = "序号"
tab.cell(5,1).text = "文件名称"
tab.cell(5,3).text = "模块总数"
tab.cell(5,4).text = "最大/最小模块行数"
line = 5
print(tmp_names)
print(function_count)
for i in range(len(tmp_names)):
    tab.cell(line+i+1, 0).text = str(i+1)
    tab.cell(line + i + 1, 1).text = tmp_names[i]
    tab.cell(line + i + 1, 3).text = function_count[i]
    tab.cell(line + i + 1, 4).text = "报告未给出"
line = line+1+len(tmp_names)
tab.cell(line,0).text = "模块圈复杂度、扇出数统计"
tab.cell(line+1,0).text = "序号"
tab.cell(line+1,1).text = "模块名称"
tab.cell(line+1,3).text = "圈复杂度（超过10）"
tab.cell(line+1,4).text = "扇出数（大于7）"
line = line+1
tmp_func = []
for i in function_list:
    for j in i:
        tmp_func.append(j)
tmp_co = []
for i in com:
    for j in i:
        tmp_co.append(j)

tmp_fan = []
for i in fan_out:
    for j in i:
        tmp_fan.append(j)

for i in range(lent):
    tab.cell(line + 1+i, 0).text = str(i+1)
    tab.cell(line + 1 + i, 1).text = tmp_func[i]
    tab.cell(line + 1 + i, 3).text = str(tmp_co[i])
    tab.cell(line + 1 + i, 4).text = str(tmp_fan[i])
line = line+1+lent
tab.cell(line,0).text = "代码结构分析"
tab.cell(line+1,0).text = "最大调用深度"
tab.cell(line+2,0).text = "数据流分析"
tab.cell(line+1,2).text = "Unknown"
tab.cell(line+2,1).text = "是否存在未引用变量再次赋值"
tab.cell(line+2,3).text = str(unused_p_assign_twice)
tab.cell(line+3,1).text = "是否存在未赋值前被引用"
tab.cell(line+3,3).text = str(used_before_assign)
tab.cell(line+4,1).text = "是否存在赋值后未引用"
tab.cell(line+4,3).text = str(unused_after_assign)

tab.cell(line+5,0).text = "控制流分析"
tab.cell(line+5,1).text = "是否含不可达语句"
tab.cell(line+5,3).text = str(unreachable_code)
tab.cell(line+6,1).text = "是否存在无限循环"
tab.cell(line+6,3).text = str(infinite_loop)
tab.cell(line+7,1).text = "是否存在未调用的函数"
tab.cell(line+7,3).text = str(uncalled_function)
tab.cell(line+8,1).text = "for循环次数是否正确"
tab.cell(line+8,3).text = str(for_cycle)

tab.cell(line+9,0).text = "接口分析"
tab.cell(line+9,1).text = "形参和实参数量是否一致、顺序是否正确"
tab.cell(line+9,3).text = str(formal_actual_p_same_count)
tab.cell(line+10,1).text = "形参和实参的属性是否一致"
tab.cell(line+10,3).text = str(formal_actual_p_same_shuxing)

tab.cell(line+11,0).text = "表达式分析"
tab.cell(line+11,1).text = "逻辑表达式中是否正确的使用括号"
tab.cell(line+11,3).text = str(brackets_error)
tab.cell(line+12,1).text = "是否存在数组下标越界"
tab.cell(line+12,3).text = str(array_bound_exceeded)
tab.cell(line+13,1).text = "是否有除零操作"
tab.cell(line+13,3).text = str(devide_zero)

line = line+13
tab.cell(line+1,0).text = "函数调用结构图"+"自己画"
tab.cell(line+2,0).text = "未调用函数分析"
tab.cell(line+2,2).text = "静态结果显示为被调用程序有：\n没说"

#填充文本

doc.save("testbed_"+software+".docx")






